import os
from abc import ABCMeta, abstractmethod
from contextlib import contextmanager
from dataclasses import dataclass, field
from typing import List, Tuple, Union, Dict, Optional
from enum import Enum, unique

import pdfplumber
from win32com import client as wc
from docx import Document


class SingletonMeta(type):

    _instances = {}

    def __call__(cls, *args, **kwargs):
        if cls not in cls._instances:
            instance = super().__call__(*args, **kwargs)
            cls._instances[cls] = instance
        return cls._instances[cls]


class SingletonABCMeta(ABCMeta):

    _instances = {}

    def __call__(cls, *args, **kwargs):
        if cls not in cls._instances:
            instance = super().__call__(*args, **kwargs)
            cls._instances[cls] = instance
        return cls._instances[cls]


@unique
class TableSearchType(Enum):
    SPLIT_LINE_BREAK_END = 1  # split by line break (\n) and check if the last element matches the keyword
    SPLIT_LINE_BREAK_START = 2  # split by line break (\n) and check if the first element matches the keyword
    REMOVE_LINE_BREAK_CONTAIN = 3  # remove all line breaks (\n) and check if contains the keyword
    SPLIT_LINE_BREAK_ALL_DIGIT = 4  # split by line break (\n) and check if all the elements are integer.
    EXACT_MATCH = 5  # check if the given string matches exactly the expected string.


@unique
class Direction(Enum):
    TRANSVERSE = 'Transverse'  # 横向
    LONGITUDINAL = 'Longitudinal'  # 纵向

    def __str__(self):
        return self._name_


@dataclass
class CertificateElement:
    table_index: Optional[int]
    x_coordinate: Optional[int]
    y_coordinate: Optional[int]
    value: str


@dataclass
class CertificateElementInPlate(CertificateElement):
    index: Optional[int]


@dataclass
class CertificateElementToVerify(CertificateElementInPlate):
    valid_flag: bool
    message: Optional[str]

    def is_valid(self):
        return self.valid_flag


@dataclass
class Specification(CertificateElementToVerify):
    pass


@dataclass
class SerialNumber(CertificateElement):
    value: int


@dataclass
class SerialNumbers(CertificateElement):
    value: List[SerialNumber]

    def __len__(self):
        return len(self.value)

    def __getitem__(self, key):
        return self.value[key]


@dataclass
class ChemicalElementName(CertificateElement):
    row_index: int
    precision: int


# BAOSHAN: Delivery Condition
# LONGTENG: Condition of Supply
@dataclass
class DeliveryCondition(CertificateElementToVerify):
    pass


# Uniform the unit to Ton.
@dataclass
class Mass(CertificateElementInPlate):
    value: float


@dataclass
class PositionDirectionImpact(CertificateElementInPlate):
    value: Direction


# BAOSHAN: HEAT NO. 炉号
# LONGTENG: BATCH NO. 冶炼炉号
@dataclass
class BatchNo(CertificateElementInPlate):
    pass


# BAOSHAN: PLATE NO. 钢板号
# LONGTENG: ROLL NO. 轧制批号
@dataclass
class PlateNo(CertificateElementInPlate):
    pass


# BAOSHAN: QTY 数量
# LONGTENG: PCS 支数
@dataclass
class Quantity(CertificateElementInPlate):
    value: int


@dataclass
class Thickness(CertificateElementToVerify):
    value: float


# Steel Making Type is required in specific rules for LongTeng
@dataclass
class SteelMakingType(CertificateElementInPlate):
    value: Optional[str]


@dataclass
class ChemicalElementValue(CertificateElementToVerify):
    value: Optional[float]
    element: str
    precision: Optional[int]
    calculated_value: Optional[float] = field(init=False)

    def __post_init__(self):
        if self.value is None or self.precision is None:
            self.calculated_value = None
        elif self.precision == 0:
            self.calculated_value = self.value
        else:
            self.calculated_value = round(self.value * (10 ** -self.precision), self.precision)

    def __str__(self):
        return (
            f"{self.element}(value={self.value}, precision={self.precision}, calculated_value={self.calculated_value})"
        )

    def set_value(self, value: float):
        self.value = value
        self.__post_init__()

    def set_precision(self, precision: int):
        self.precision = precision
        self.__post_init__()

    def set_value_and_precision(self, value: float, precision: int):
        self.value = value
        self.precision = precision
        self.__post_init__()


@dataclass
class YieldStrength(CertificateElementToVerify):
    value: int


@dataclass
class TensileStrength(CertificateElementToVerify):
    value: int


@dataclass
class Elongation(CertificateElementToVerify):
    value: int


@dataclass
class Temperature(CertificateElementToVerify):
    value: int


@dataclass
class ImpactEnergy(CertificateElementToVerify):
    value: int
    test_number: str


class SteelPlate:

    def __init__(self, serial_number: SerialNumber):
        self.serial_number: SerialNumber = serial_number
        self.batch_no: Optional[BatchNo] = None
        self.plate_no: Optional[PlateNo] = None
        self.specification: Optional[Specification] = None
        self.thickness: Optional[Thickness] = None
        self.quantity: Optional[Quantity] = None
        self.mass: Optional[Mass] = None
        self.chemical_compositions: Dict[str, ChemicalElementValue] = dict()
        self.yield_strength: Optional[YieldStrength] = None
        self.tensile_strength: Optional[TensileStrength] = None
        self.elongation: Optional[Elongation] = None
        self.position_direction_impact: Optional[PositionDirectionImpact] = None
        self.temperature: Optional[Temperature] = None
        self.impact_energy_list: List[ImpactEnergy] = []
        self.delivery_condition: Optional[DeliveryCondition] = None
        # LongTeng
        self.steel_making_type: Optional[SteelMakingType] = None

    def __str__(self):
        chemical_str = '\n\t\t\t'.join(
            [str(self.chemical_compositions[element]) for element in
             self.chemical_compositions])
        position_direction_str = \
            'None' if self.position_direction_impact is None else self.position_direction_impact.value
        temperature_str = 'None' if self.temperature is None else self.temperature.value
        impact_energy_str = ', '.join(
            ['None' if impact_energy is None else f"{impact_energy.test_number}: {str(impact_energy.value)}" for
             impact_energy in self.impact_energy_list])
        return (
            # f"\tSteel Plate:\n"
            f"\t\tSerial Number: {self.serial_number}\n"
            f"\t\tPlate No.: {self.plate_no}\n"
            f"\t\tMass: {self.mass.value}\n"
            f"\t\tChemical Composition:\n"
            f"\t\t\t{chemical_str}\n"
            f"\t\tYield Strength: {self.yield_strength.value}\n"
            f"\t\tTensile Strength: {self.tensile_strength.value}\n"
            f"\t\tElongation: {self.elongation.value}\n"
            f"\t\tPosition Direction of Impact Test: {position_direction_str}\n"
            f"\t\tTemperature: {temperature_str}\n"
            f"\t\tAbsorbed Energy: {impact_energy_str}\n"
            f"\t\tDelivery Condition: {self.delivery_condition.value}\n"
        )


@dataclass
class Certificate(metaclass=ABCMeta):
    file_path: str
    steel_plant: str
    certificate_no: str
    # specification: Specification
    # thickness: Thickness
    serial_numbers: Optional[SerialNumbers]
    steel_plates: Optional[List[SteelPlate]]
    chemical_elements: Optional[Dict[str, ChemicalElementName]]

    def __str__(self):
        chemical_element_str = {element: self.chemical_elements[element].precision for element in
                                self.chemical_elements}
        steel_plates_str = '\n'.join([str(steel_plate) for steel_plate in self.steel_plates])
        return (
            f"BaoSteelCertificate:\n"
            f"\tFile Path: {self.file_path}\n"
            f"\tSteel Plant: {self.steel_plant}\n"
            # f"\tSpecification: {self.specification.value}\n"
            # f"\tThickness: {self.thickness.value}\n"
            f"\tChemical Elements: {chemical_element_str}\n"
            f"\tSerial Numbers: {self.serial_numbers.value}\n"
            f"\tSteel Plates:\n{steel_plates_str}"
        )


class Limit(metaclass=ABCMeta):

    @abstractmethod
    def verify_value(self, value) -> Tuple[bool, str]:
        pass

    @abstractmethod
    def verify(self, plate: SteelPlate) -> bool:
        pass

    @abstractmethod
    # def get_element(self, certificate: Certificate, steel_plate_index: int):
    def get_element(self, plate: SteelPlate):
        pass


class CertificateFile:

    def __init__(self, file_path: str):
        self.file_path = file_path if os.path.isabs(file_path) else os.path.abspath(file_path)


class PdfFile(CertificateFile):

    def __enter__(self):
        self.pdf = pdfplumber.open(self.file_path)
        self.page = self.pdf.pages[0]  # Always has only one page
        self.tables = self.page.extract_tables()
        self.content = self.page.extract_text()
        self.steel_plant = self.extract_steel_plant()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        # pass
        if self.pdf:
            self.pdf.close()

    def extract_steel_plant(self):
        if "No.885 Fujin ROAD, BAOSHAN DISTRICT".replace(" ", "").upper() in self.content.replace(" ", "").upper():
            return 'BAOSHAN IRON & STEEL CO., LTD.'
        else:
            raise ValueError(
                f"The steel plant name could not be recognized for the given PDF file {self.file_path}."
            )


class DocxFile(CertificateFile):

    def __enter__(self):
        ext = os.path.splitext(self.file_path)[-1]
        if ext == '.docx':
            pass
        elif ext == '.doc':
            # Convert doc format to docx
            word = wc.Dispatch("Word.Application")
            doc = word.Documents.Open(self.file_path)
            doc.SaveAs(f"{self.file_path}x", 12)
            doc.Close()
            word.Quit()
            os.remove(self.file_path)
            self.file_path = f"{self.file_path}x"
        else:
            raise ValueError(
                f"The extension name of file path {self.file_path} passed to DocxFile constructor is neither docx "
                f"nor doc."
            )
        self.document = Document(self.file_path)
        self.steel_plant = self.extract_steel_plant()
        self.tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in self.document.tables]
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass

    def extract_steel_plant(self):
        steel_plant = None
        for para in self.document.paragraphs:
            if 'CHANGSHU LONGTENG SPECIAL STEEL CO., LTD' in para.text:
                steel_plant = 'CHANGSHU LONGTENG SPECIAL STEEL CO., LTD.'
                break
        if steel_plant is not None:
            return steel_plant
        else:
            raise ValueError(
                f"The steel plant name could not be recognized in the given file {self.file_path}"
            )


class CommonUtils:

    chemical_elements_table = [
        'C', 'Si', 'Mn', 'P', 'S', 'Cu', 'Cr', 'Ni', 'Mo', 'Ceq', 'Als', 'Alt', 'Nb', 'Ti', 'V', 'Al'
    ]

    @staticmethod
    @contextmanager
    def open_file(file_path: str) -> CertificateFile:
        if file_path.lower().endswith('.pdf'):
            with PdfFile(file_path) as pdf_file:
                try:
                    yield pdf_file
                finally:
                    # pdf_file.pdf.close()
                    pass
        elif file_path.lower().endswith('.doc') or file_path.lower().endswith('.docx'):
            with DocxFile(file_path) as docx_file:
                try:
                    yield docx_file
                finally:
                    pass
        else:
            raise ValueError(
                f"The extension of the given file_path {file_path} is not valid, failed to open file."
            )

    # Ordinal numbers replacement
    @staticmethod
    def ordinal(n: int):
        return "%d%s" % (n, "tsnrhtdd"[(n // 10 % 10 != 1) * (n % 10 < 4) * n % 10::4])

    @staticmethod
    def search_table(
        table: List[List[Union[str, None]]],
        keyword: Union[str, None],
        search_type: TableSearchType = TableSearchType.SPLIT_LINE_BREAK_END,
        confirmed_row: int = None,
        confirmed_col: int = None
    ) -> Union[Tuple[int, int], None]:

        # Define search methods:
        search_methods = {
            TableSearchType.SPLIT_LINE_BREAK_END:
                lambda table_cell: table_cell is not None and table_cell.split('\n')[-1].strip() == keyword,
            TableSearchType.SPLIT_LINE_BREAK_START:
                lambda table_cell: table_cell is not None and table_cell.split('\n')[0].strip() == keyword,
            TableSearchType.REMOVE_LINE_BREAK_CONTAIN:
                lambda table_cell: table_cell is not None and keyword in table_cell.replace('\n', '').replace(' ', ''),
            TableSearchType.SPLIT_LINE_BREAK_ALL_DIGIT:
                lambda table_cell: table_cell is not None and all(
                    map(lambda x: x.strip().isdigit(), table_cell.split('\n'))),
            TableSearchType.EXACT_MATCH:
                lambda table_cell: table_cell is not None and table_cell == keyword
        }

        coordinates = None

        if confirmed_row is None:
            if confirmed_col is None:
                for row_index, row in enumerate(table):
                    if coordinates is not None:
                        break
                    for col_index, cell in enumerate(row):
                        if search_methods[search_type](cell):
                            coordinates = (row_index, col_index)
                            break
            else:
                for row_index, row in enumerate(table):
                    cell = row[confirmed_col]
                    if search_methods[search_type](cell):
                        coordinates = (row_index, confirmed_col)
                        break
        else:
            if confirmed_col is None:
                row = table[confirmed_row]
                for col_index, cell in enumerate(row):
                    if search_methods[search_type](cell):
                        coordinates = (confirmed_row, col_index)
                        break
            else:
                cell = table[confirmed_row][confirmed_col]
                if search_methods[search_type](cell):
                    coordinates = (confirmed_row, confirmed_col)

        return coordinates

    @staticmethod
    def verify_chemical_element_limit(element: str, chemical_composition_limit: dict, element_calculated_value: float):
        if chemical_composition_limit['type'] == 'maximum':
            if element_calculated_value <= chemical_composition_limit['limit']:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, meets "
                    f"the maximum limit {chemical_composition_limit['limit']}."
                )
            else:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, violates "
                    f"the maximum limit {chemical_composition_limit['limit']}."
                )
                return False
        elif chemical_composition_limit['type'] == 'minimum':
            if element_calculated_value >= chemical_composition_limit['limit']:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, meets "
                    f"the minimum limit {chemical_composition_limit['limit']}."
                )
            else:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, violates "
                    f"the minimum limit {chemical_composition_limit['limit']}."
                )
                return False
        elif chemical_composition_limit['type'] == 'range':
            if chemical_composition_limit['minimum'] <= element_calculated_value <= \
                    chemical_composition_limit['maximum']:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, meets "
                    f"the valid range [{chemical_composition_limit['minimum']}, "
                    f"{chemical_composition_limit['maximum']}]."
                )
            else:
                print(
                    f"The value of chemical element {element} is {element_calculated_value}, violates "
                    f"the valid range [{chemical_composition_limit['minimum']}, "
                    f"{chemical_composition_limit['maximum']}]."
                )
                return False
        else:
            raise ValueError(
                f"The chemical composition limit type {chemical_composition_limit['type']} of "
                f"chemical element {element} is invalid!"
            )
        return True
